import docx
import io

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from a DOCX file byte stream.
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        return "\n".join(full_text)
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return ""
